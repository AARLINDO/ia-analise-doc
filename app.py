import streamlit as st
import streamlit.components.v1 as components
import os
import json
import time
import re
import random
from datetime import datetime, date
from io import BytesIO

# =============================================================================
# 1. CONFIGURAÇÃO INICIAL
# =============================================================================
st.set_page_config(
    page_title="Carmélio AI | Ultimate Studio",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =============================================================================
# 2. IMPORTAÇÕES E SETUP
# =============================================================================
try: 
    import google.generativeai as genai
except ImportError: 
    genai = None

try: 
    import pdfplumber
except ImportError: 
    pdfplumber = None

try: 
    import docx
    from docx import Document
except ImportError: 
    docx = None
    Document = None

try: 
    from PIL import Image
except ImportError: 
    Image = None

# Inicialização de Estado (Session State)
keys = {
    "user_xp": 0, "contract_step": 1, "contract_clauses": [], 
    "contract_meta": {}, "chat_history": [], "edital_text": "", 
    "edital_filename": "", "quiz_data": None, "quiz_show_answer": False, 
    "user_choice": None, "ocr_text": "", "last_call": 0, "audio_text": ""
}
for k, v in keys.items():
    if k not in st.session_state: 
        st.session_state[k] = v

# =============================================================================
# 3. FUNÇÕES UTILITÁRIAS E LÓGICA (BACKEND)
# =============================================================================

def check_rate_limit():
    """Evita chamadas excessivas (proteção simples)."""
    if time.time() - st.session_state.last_call < 2.0: 
        return True 
    return False

def mark_call(): 
    st.session_state.last_call = time.time()

def add_xp(amount):
    st.session_state.user_xp += amount
    st.toast(f"+{amount} XP | Nível {int(st.session_state.user_xp/100)}", icon="⚡")

@st.cache_resource
def get_best_model():
    """Configura e retorna o melhor modelo Gemini disponível (Modelos Top Atualizados)."""
    api_key = st.secrets.get("GOOGLE_API_KEY")
    if not api_key: 
        return None, "⚠️ Configure secrets.toml"
    try:
        genai.configure(api_key=api_key)
        try: 
            models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        except: 
            return None, "Erro de Chave API"
        
        pref = [
            'models/gemini-1.5-pro',         
            'models/gemini-1.5-flash',       
            'models/gemini-1.5-flash-latest',
            'models/gemini-pro'
        ]
        escolhido = next((m for m in pref if m in models), models[0] if models else None)
        
        if escolhido: 
            return genai.GenerativeModel(escolhido.replace("models/", "")), escolhido.replace("models/", "")
        return None, "Nenhum modelo compatível."
    except Exception as e: 
        return None, f"Erro Fatal: {str(e)}"

def call_gemini(system_prompt, user_prompt, json_mode=False, image=None, audio_bytes=None, audio_mime=None, use_search=False):
    """Função central de comunicação com a IA com suporte a texto, imagem e áudio nativo."""
    if check_rate_limit(): 
        time.sleep(1)
    
    mark_call()
    model, name = get_best_model()
    if not model: 
        return f"Erro: {name}"
    
    try:
        tools_config = 'google_search_retrieval' if use_search else None
        
        # Fluxo Multimodal para Áudio Real
        if audio_bytes:
            audio_part = {
                "mime_type": audio_mime,
                "data": audio_bytes
            }
            response = model.generate_content([system_prompt, audio_part, user_prompt])
            return response.text

        # Fluxo Multimodal para Imagem (OCR)
        if image:
            response = model.generate_content([system_prompt, image, user_prompt])
            return response.text
            
        # Fluxo Tradicional de Texto
        full_prompt = f"SYSTEM ROLE: {system_prompt}\nUSER REQUEST: {user_prompt}"
        if json_mode: 
            full_prompt += "\nFORMAT: Return ONLY valid JSON. No Markdown."
        
        if tools_config:
            try:
                response = model.generate_content(full_prompt, tools=tools_config)
            except:
                response = model.generate_content(full_prompt)
        else:
            response = model.generate_content(full_prompt)
                
        return response.text
    except Exception as e: 
        if "429" in str(e):
            return "⚠️ Limite de velocidade atingido. Aguarde 30 segundos e tente novamente."
        return f"Erro IA: {str(e)}"

def extract_json_surgical(text):
    """Extrai JSON de texto bagunçado."""
    try:
        text = text.replace("```json", "").replace("```", "")
        match = re.search(r"(\{[\s\S]*\}|\[[\s\S]*\])", text)
        if match: 
            return json.loads(match.group(0))
    except: 
        pass
    return None

def read_pdf_safe(file_obj):
    """Lê PDF e retorna texto."""
    if not pdfplumber: 
        return None
    try:
        text = ""
        with pdfplumber.open(BytesIO(file_obj.getvalue())) as pdf:
            for i, p in enumerate(pdf.pages):
                if i >= 300: 
                    break 
                text += (p.extract_text() or "") + "\n"
        return text if text.strip() else None
    except: 
        pass
    return None

def create_generic_docx(content, title="Documento Carmélio AI"):
    if not docx: 
        return None
    doc = Document()
    doc.add_heading(title, 0)
    for line in content.split('\n'):
        if line.strip(): 
            doc.add_paragraph(line.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_contract_docx(clauses, meta):
    if not docx: 
        return None
    doc = Document()
    doc.add_heading(meta.get('tipo', 'CONTRATO').upper(), 0)
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_heading("1. QUALIFICAÇÃO", level=1)
    doc.add_paragraph(meta.get('partes', ''))
    doc.add_heading("2. DO OBJETO", level=1)
    doc.add_paragraph(meta.get('objeto', ''))
    for clause in clauses:
        doc.add_heading(clause.get('titulo', 'Cláusula'), level=1)
        for line in clause.get('conteudo', '').split('\n'):
            if line.strip(): 
                doc.add_paragraph(line.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def get_daily_verse():
    """Retorna um versículo dinâmico e diferente baseado rigorosamente no dia do ano."""
    versiculos = [
        {"ref": "Josué 1:9", "txt": "Seja forte e corajoso! Não se apavore nem desanime, pois o Senhor, o seu Deus, estará com você por onde você andar."},
        {"ref": "Filipenses 4:13", "txt": "Tudo posso naquele que me fortalece."},
        {"ref": "Salmos 37:5", "txt": "Entregue o seu caminho ao Senhor; confie nele, e ele agirá."},
        {"ref": "Isaías 41:10", "txt": "Não tema, pois estou com você; não tenha medo, pois sou o seu Deus. Eu o fortalecerei e o ajudarei."},
        {"ref": "Jeremias 29:11", "txt": "Porque sou eu que conheço os planos que tenho para vocês, diz o Senhor, planos de fazê-los prosperar."},
        {"ref": "Provérbios 16:3", "txt": "Consagre ao Senhor tudo o que você faz, e os seus planos serão bem-sucedidos."},
        {"ref": "Salmos 121:1-2", "txt": "Levanto os meus olhos para os montes e pergunto: De onde me vem o socorro? O meu socorro vem do Senhor."},
        {"ref": "2 Timóteo 1:7", "txt": "Pois Deus não nos deu espírito de covardia, mas de poder, de amor e de equilíbrio."},
        {"ref": "Salmos 23:1", "txt": "O Senhor é o meu pastor; de nada terei falta."},
        {"ref": "Isaías 40:31", "txt": "Mas aqueles que esperam no Senhor renovam as suas forças. Voam bem alto como águias; correm e não ficam exaustos."},
        {"ref": "Mateus 6:33", "txt": "Busquem, pois, em primeiro lugar o Reino de Deus e a sua justiça, e todas essas coisas serão acrescentadas a vocês."},
        {"ref": "Salmos 46:1", "txt": "Deus é o nosso refúgio e a nossa fortaleza, auxílio sempre presente na adversidade."},
        {"ref": "Romanos 8:28", "txt": "Sabemos que Deus age em todas as coisas para o bem daqueles que o amam, dos que foram chamados de acordo com o seu propósito."},
        {"ref": "Provérbios 3:5", "txt": "Confie no Senhor de todo o seu coração e não se apóie em seu próprio entendimento."},
        {"ref": "Salmos 119:105", "txt": "A tua palavra é lâmpada que ilumina os meus passos e luz que clareia o meu caminho."},
        {"ref": "João 16:33", "txt": "No mundo vocês terão aflições; contudo, tenham ânimo! Eu venci o mundo."},
        {"ref": "Gálatas 6:9", "txt": "E não nos cansemos de fazer o bem, pois no tempo próprio colheremos, se não desanimarmos."},
        {"ref": "Salmos 27:1", "txt": "O Senhor é a minha luz e a minha salvação; de quem terei medo? O Senhor é a fortaleza da minha vida."},
        {"ref": "Tiago 1:5", "txt": "Se algum de vocês tem falta de sabedoria, peça-a a Deus, que a todos dá livremente, de boa vontade."},
        {"ref": "Romanos 12:12", "txt": "Alegrem-se na esperança, sejam pacientes na tribulação, perseverem na oração."},
        {"ref": "1 Coríntios 16:14", "txt": "Façam tudo com amor."},
        {"ref": "Salmos 34:17", "txt": "Os justos clamam, o Senhor os ouve e os livra de todas as suas tribulações."},
        {"ref": "Hebreus 11:1", "txt": "Ora, a fé é a certeza de que haverá de receber o que se espera, e a prova das coisas que não se vêem."},
        {"ref": "Salmos 91:1", "txt": "Aquele que habita no abrigo do Altíssimo e descansa à sombra do Todo-Poderoso pode dizer ao Senhor: Tu és o meu refúgio."},
        {"ref": "Deuteronômio 31:6", "txt": "Sejam fortes e corajosos. Não tenham medo nem fiquem apavorados por causa deles, pois o Senhor, o seu Deus, vai com vocês."},
        {"ref": "Colossenses 3:23", "txt": "Tudo o que fizerem, façam de todo o coração, como para o Senhor, e não para os homens."},
        {"ref": "Salmos 118:24", "txt": "Este é o dia que o Senhor fez; exultemos e alegremo-nos nele."},
        {"ref": "1 Pedro 5:7", "txt": "Lancem sobre ele toda a sua ansiedade, porque ele tem cuidado de vocês."},
        {"ref": "Provérbios 4:23", "txt": "Acima de tudo o que deve ser preservado, guarde o seu coração, porque dele procedem as fontes da vida."},
        {"ref": "Efésios 6:10", "txt": "Finalmente, fortaleçam-se no Senhor e no seu forte poder."},
        {"ref": "Salmos 139:14", "txt": "Eu te louvo porque me fizeste de modo especial e admirável. Tuas obras são maravilhosas!"}
    ]
    dia_do_ano = date.today().timetuple().tm_yday
    index = dia_do_ano % len(versiculos)
    return versiculos[index]

# =============================================================================
# 4. INTERFACE GRÁFICA & CSS
# =============================================================================
def safe_image_show(image_path):
    if os.path.exists(image_path):
        try: 
            st.image(image_path, use_container_width=True)
        except TypeError: 
            st.image(image_path, use_column_width=True)
    else: 
        st.markdown("## ⚖️ Carmélio AI")

st.markdown("""
<style>
    .stApp { background-color: #0E1117; color: #E0E0E0; }
    [data-testid="stSidebar"] { background-color: #11141d; border-right: 1px solid #2B2F3B; }
    .gemini-text {
        background: -webkit-linear-gradient(45deg, #4285F4, #9B72CB, #D96570);
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        font-weight: 800; font-size: 2.2rem;
    }
    div.stButton > button {
        background: linear-gradient(90deg, #2563EB 0%, #7C3AED 100%);
        color: white; border: none; font-weight: 600; border-radius: 8px;
    }
    .footer-credits { 
        text-align: center; margin-top: 30px; padding-top: 20px;
        border-top: 1px solid #2B2F3B; color: #6B7280; font-size: 11px; 
    }
    .footer-credits strong { color: #E0E0E0; }
    .onboarding-box {
        background-color: #1F2430; padding: 20px; border-radius: 10px;
        border-left: 5px solid #4285F4; margin-bottom: 20px;
    }
</style>
""", unsafe_allow_html=True)

def render_sidebar_widgets():
    v = get_daily_verse()

    html_code = f"""
    <style>
        .widget-box {{
            background: #1F2430;
            border: 1px solid #374151;
            border-radius: 12px;
            padding: 12px;
            text-align: center;
            color: white;
            font-family: sans-serif;
            margin-bottom: 12px;
        }}

        .devotional-box {{
            background: linear-gradient(135deg, #1e293b, #0f172a);
            border-left: 4px solid #F59E0B;
            text-align: left;
            padding: 12px;
            margin-bottom: 15px;
            border-radius: 8px;
        }}

        .verse-text {{
            font-style: italic;
            font-size: 12px;
            color: #E2E8F0;
            margin-bottom: 5px;
        }}

        .verse-ref {{
            font-size: 10px;
            font-weight: bold;
            color: #F59E0B;
            text-align: right;
        }}

        .time-display {{
            font-size: 34px;
            font-weight: 800;
            margin: 6px 0;
            color: #4285F4;
            transition: color 0.3s ease;
        }}

        .config-label {{
            font-size: 10px;
            color: #8B949E;
            font-weight: bold;
            margin-top: 6px;
            margin-bottom: 3px;
            text-align: left;
            padding-left: 4px;
        }}

        .selector-group {{
            display: flex;
            justify-content: space-between;
            background: #11141d;
            padding: 2px;
            border-radius: 6px;
            margin-bottom: 8px;
        }}

        .selector-btn {{
            flex: 1;
            background: transparent;
            border: none;
            color: #8B949E;
            padding: 4px 0;
            font-size: 11px;
            font-weight: 600;
            cursor: pointer;
            border-radius: 4px;
            transition: all 0.2s;
        }}

        .selector-btn.active {{
            background: #2563EB;
            color: white;
        }}

        .btn {{
            border: none;
            padding: 6px 12px;
            border-radius: 6px;
            cursor: pointer;
            margin: 2px;
            font-size: 12px;
            font-weight: bold;
            color: white;
        }}

        .btn-primary {{ background: #2563EB; }}
        .btn-warn {{ background: #D97706; }}
        .btn-danger {{ background: #DC2626; }}

        audio {{
            width: 100%;
            margin-top: 10px;
        }}
    </style>

    <div class="devotional-box">
        <div style="color:#F59E0B;font-size:11px;font-weight:bold;margin-bottom:4px;">
            <table>📖 Palavra do Dia</table>
        </div>
        <div class="verse-text">
            "{v['txt']}"
        </div>
        <div class="verse-ref">
            {v['ref']}
        </div>
    </div>

    <div class="widget-box">
        <div style="font-size:11px;font-weight:bold;color:#8B949E;margin-bottom:2px;">
            TOMATO FOCUS
        </div>

        <!-- Seletores de Configuração de Tempo -->
        <div class="config-label">TEMPO DE FOCO</div>
        <div class="selector-group">
            <button class="selector-btn" id="foco-25" onclick="setConfigFoco(25)">25m</button>
            <button class="selector-btn" id="foco-35" onclick="setConfigFoco(35)">35m</button>
            <button class="selector-btn" id="foco-45" onclick="setConfigFoco(45)">45m</button>
        </div>

        <div class="config-label">INTERVALO</div>
        <div class="selector-group">
            <button class="selector-btn" id="break-5" onclick="setConfigBreak(5)">5m</button>
            <button class="selector-btn" id="break-10" onclick="setConfigBreak(10)">10m</button>
            <button class="selector-btn" id="break-15" onclick="setConfigBreak(15)">15m</button>
        </div>

        <div class="time-display" id="timer">25:00</div>

        <div id="status" style="font-size:10px;color:#aaa;margin-bottom:8px;">
             Pronto
        </div>

        <div>
            <button class="btn btn-primary" onclick="startTimer()">▶</button>
            <button class="btn btn-warn" onclick="pauseTimer()">⏸</button>
            <button class="btn btn-danger" onclick="resetTimer()">↺</button>
        </div>
    </div>

    <div class="widget-box">
        <div style="font-size:11px;font-weight:bold;color:#8B949E;margin-bottom:10px;">
            🎵 Rádio LoFi 24h
        </div>

        <audio controls>
            <source src="https://stream.zeno.fm/f3wvbbqmdg8uv" type="audio/mpeg">
            Seu navegador não suporta áudio.
        </audio>

        <div style="font-size:10px;color:#34D399;margin-top:8px;">
            Clique em Play para ouvir
        </div>
    </div>

    <script>
        let interval = null;

        // Carrega configurações personalizadas ou adota os padrões iniciais (25/5)
        let cfgFoco = localStorage.getItem('pomodoro_cfg_foco') ? parseInt(localStorage.getItem('pomodoro_cfg_foco')) : 25;
        let cfgBreak = localStorage.getItem('pomodoro_cfg_break') ? parseInt(localStorage.getItem('pomodoro_cfg_break')) : 5;

        let time = localStorage.getItem('pomodoro_time') ? parseInt(localStorage.getItem('pomodoro_time')) : (cfgFoco * 60);
        let isModeFoco = localStorage.getItem('pomodoro_mode') === 'false' ? false : true;
        let isRunning = localStorage.getItem('pomodoro_running') === 'true' ? true : false;

        function setConfigFoco(mins) {{
            if(isRunning) return; // Trava alteração com o timer rodando
            cfgFoco = mins;
            localStorage.setItem('pomodoro_cfg_foco', mins);
            if(isModeFoco) {{
                time = mins * 60;
                localStorage.setItem('pomodoro_time', time);
            }}
            updateUI();
        }}

        function setConfigBreak(mins) {{
            if(isRunning) return;
            cfgBreak = mins;
            localStorage.setItem('pomodoro_cfg_break', mins);
            if(!isModeFoco) {{
                time = mins * 60;
                localStorage.setItem('pomodoro_time', time);
            }}
            updateUI();
        }}

        function updateUI() {{
            // Atualiza o relógio
            let m = Math.floor(time / 60);
            let s = time % 60;
            document.getElementById('timer').innerText = (m < 10 ? '0' : '') + m + ':' + (s < 10 ? '0' : '') + s;

            // Altera cores de acordo com o estado ativo
            document.getElementById('timer').style.color = isModeFoco ? '#4285F4' : '#34D399';
            document.getElementById('status').innerText = isRunning ? (isModeFoco ? 'Focando...' : 'Descansando...') : 'Pausado';

            // Altera as classes ativas dos botões de configuração
            ['25','35','45'].forEach(v => {{
                document.getElementById('foco-'+v).classList.remove('active');
            }});
            document.getElementById('foco-'+cfgFoco).classList.add('active');

            ['5','10','15'].forEach(v => {{
                document.getElementById('break-'+v).classList.remove('active');
            }});
            document.getElementById('break-'+cfgBreak).classList.add('active');
        }}

        function startTimer() {{
            if(interval) return;
            
            isRunning = true;
            localStorage.setItem('pomodoro_running', 'true');

            interval = setInterval(() => {{
                if(time > 0) {{
                    time--;
                    localStorage.setItem('pomodoro_time', time);
                    updateUI();
                }}
                else {{
                    clearInterval(interval);
                    interval = null;
                    
                    if(isModeFoco) {{
                        isModeFoco = false;
                        time = cfgBreak * 60;
                        localStorage.setItem('pomodoro_mode', 'false');
                        alert('Ciclo de Foco Concluído! Hora do seu intervalo de ' + cfgBreak + ' minutos.');
                    }} else {{
                        isModeFoco = true;
                        time = cfgFoco * 60;
                        localStorage.setItem('pomodoro_mode', 'true');
                        alert('Fim do descanso! Iniciando bloco de foco de ' + cfgFoco + ' minutos.');
                    }}
                    
                    localStorage.setItem('pomodoro_time', time);
                    updateUI();
                    startTimer();
                }}
            }}, 1000);
        }}

        function pauseTimer() {{
            clearInterval(interval);
            interval = null;
            isRunning = false;
            localStorage.setItem('pomodoro_running', 'false');
            updateUI();
        }}

        function resetTimer() {{
            pauseTimer();
            isModeFoco = true;
            time = cfgFoco * 60;
            localStorage.setItem('pomodoro_time', time);
            localStorage.setItem('pomodoro_mode', 'true');
            localStorage.setItem('pomodoro_running', 'false');
            updateUI();
            document.getElementById('status').innerText = 'Pronto';
        }}

        // Execução inicial
        updateUI();
        if (isRunning) {{
            startTimer();
        }}
    </script>
    """
    components.html(html_code, height=600)
# =============================================================================
# 5. EXECUÇÃO PRINCIPAL E FLUXO DE TELAS (ATUALIZADO)
# =============================================================================
with st.sidebar:
    safe_image_show("carmelio_logo.png.png")
    render_sidebar_widgets()
    st.markdown("---")
    
    # Sistema Gatekeeper (Agora protege apenas os módulos Premium)
    password_input = st.text_input("Chave do Estúdio (Premium):", type="password", placeholder="Insira a senha master")
    studio_authorized = (password_input == st.secrets.get("STUDIO_PASSWORD", "1234"))
    
    if not studio_authorized:
        st.warning("🔒 Insira a Chave para liberar as ferramentas Premium.")
    
    model_obj, status_msg = get_best_model()
    if not model_obj: 
        st.error(f"❌ {status_msg}")
    else: 
        st.success(f"🟢 **Modelo Ativo: {status_msg}**")
        
    # Menu reorganizado: OAB agora é a primeira opção (Tela de Entrada)
    menu = st.radio("Menu", [
        "🎓 Gabaritando a OAB", 
        "✨ Chat Inteligente", 
        "📝 Gere seu Contrato", 
        "🎯 Mestre dos Editais", 
        "🏢 Cartório OCR", 
        "🎙️ Transcrição"
    ], label_visibility="collapsed")
    
    st.markdown("---")
    st.progress(min((st.session_state.user_xp % 100) / 100, 1.0))
    st.markdown("""<div class='footer-credits'>Desenvolvido por<br><strong>Arthur Carmélio</strong><br>© 2026 Carmélio AI</div>""", unsafe_allow_html=True)


# 🛑 FUNÇÃO AUXILIAR DE BLOQUEIO PREMIUM
def verificar_acesso_premium():
    if not studio_authorized:
        st.info("### 🔐 Recursos Avançados Restritos\nEste módulo faz parte do ecossistema premium do Carmélio AI. Para liberá-lo, insira a **Chave do Estúdio** na barra lateral esquerda.")
        return False
    return True


# =============================================================================
# 6. RENDERIZAÇÃO DAS TELAS
# =============================================================================

# 🔓 MÓDULO OAB: 100% ABERTO E GRATUITO PARA QUALQUER VISITANTE
if menu == "🎓 Gabaritando a OAB":
    st.title("🎓 Gabaritando a OAB (1ª Fase)")
    st.markdown("""
    <div class="onboarding-box">
        <h4>⚡ Questões Reais da FGV em Tempo Real</h4>
        <p>Treine com o histórico oficial da OAB sem precisar de uploads. Nossa IA localiza, filtra e fundamenta as questões dinamicamente.</p>
    </div>
    """, unsafe_allow_html=True)

    materias_oab = [
        "Simulado Geral (Todas as Matérias)",
        "Ética Profissional", "Direito Constitucional", "Direito Administrativo", 
        "Direito Penal", "Processo Penal", "Direito Civil", "Processo Civil",
        "Direito do Trabalho", "Processo do Trabalho", "Direito Tributário", 
        "Direito Empresarial", "Direitos Humanos", "Direito Internacional"
    ]

    def gerar_questao_oab(materia_selecionada):
        st.session_state.oab_quiz_data = None
        st.session_state.oab_show_answer = False
        with st.spinner("🔍 Varrendo a web em busca de questões oficiais FGV..."):
            filtro_materia = "" if "Geral" in materia_selecionada else f"especificamente da matéria de {materia_selecionada}"
            
            prompt = f"""
            ROLE: Professor Especialista em Exame de Ordem da OAB.
            TASK: Use a busca da internet para encontrar uma QUESTÃO REAL E OFICIAL de exames passados da OAB aplicados pela banca FGV, {filtro_materia}.
            REQUIREMENT: Forneça o enunciado completo, as 4 alternativas oficiais e o gabarito correto.
            JSON Output: {{'exame':'Exame OAB número XXX','materia':'...','enunciado':'...','alternativas':{{'A':'...','B':'...','C':'...','D':'...'}},'correta':'A','explicacao':'...'}}
            """
            
            res = call_gemini("JSON Only.", prompt, json_mode=True, use_search=True)
            if "Limite de velocidade" in res:
                st.error(res)
            else:
                data = extract_json_surgical(res)
                if data: 
                    st.session_state.oab_quiz_data = data
                else: 
                    st.error("Erro ao estruturar a questão. Tente buscar novamente.")

    c1, c2 = st.columns([2, 1])
    with c1:
        mat_escolhida = st.selectbox("Escolha a disciplina para treinar:", materias_oab)
    with c2:
        st.write(""); st.write("")
        if st.button("🚀 TRAZER QUESTÃO", type="primary", use_container_width=True):
            gerar_questao_oab(mat_escolhida)
            st.rerun()
    if st.session_state.oab_quiz_data:
        q = st.session_state.oab_quiz_data
        st.markdown(f"### 📝 {q.get('exame', 'Exame de Ordem')} | Disciplina: {q.get('materia', mat_escolhida)}")
        st.info(q['enunciado'])
        opts = q['alternativas']
        
        if not st.session_state.oab_show_answer:
            c1, c2 = st.columns(2)
            if c1.button(f"A) {opts['A']}", use_container_width=True, key="oab_a"): 
                st.session_state.oab_choice = "A"; st.session_state.oab_show_answer = True; st.rerun()
            if c2.button(f"B) {opts['B']}", use_container_width=True, key="oab_b"): 
                st.session_state.oab_choice = "B"; st.session_state.oab_show_answer = True; st.rerun()
            if c1.button(f"C) {opts['C']}", use_container_width=True, key="oab_c"): 
                st.session_state.oab_choice = "C"; st.session_state.oab_show_answer = True; st.rerun()
            if c2.button(f"D) {opts['D']}", use_container_width=True, key="oab_d"): 
                st.session_state.oab_choice = "D"; st.session_state.oab_show_answer = True; st.rerun()
        else:
            u, c = st.session_state.oab_choice, q['correta']
            for l, t in opts.items():
                icon = "✅" if l == c else ("❌" if l == u else "⬜")
                st.write(f"{icon} **{l})** {t}")
            
            if u == c: 
                st.success("🎯 Sensacional! Você acertou uma questão oficial da OAB!")
                add_xp(60)
            else: 
                st.error(f"Faltou pouco! A alternativa correta era a letra: {c}")
            
            st.markdown(f"**📚 Justificativa e Fundamentação Legal:**")
            st.write(q['explicacao'])
            
            q_text = f"EXAME: {q.get('exame','OAB FGV')}\nMATÉRIA: {q.get('materia', mat_escolhida)}\n\nQUESTÃO:\n{q['enunciado']}\n\nA) {opts['A']}\nB) {opts['B']}\nC) {opts['C']}\nD) {opts['D']}\n\nRESPOSTA: {q['correta']}\n\nJUSTIFICATIVA:\n{q['explicacao']}"
            docx_q = create_generic_docx(q_text, "Questão Oficial OAB - Carmélio AI")
            if docx_q:
                st.download_button("💾 Baixar Questão (Word)", docx_q, "Questao_OAB_Carmelio.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            st.write("")
            if st.button("➡️ Próxima Questão", type="primary"):
                gerar_questao_oab(mat_escolhida)
                st.rerun()


# 🔒 MÓDULOS DE ACESSO RESTRITO (EXIGEM A SENHA MASTER)
elif menu == "✨ Chat Inteligente":
    if verificar_acesso_premium():
        st.markdown('<h1 class="gemini-text">Mentor Jurídico</h1>', unsafe_allow_html=True)
        if not st.session_state.chat_history: 
            st.markdown("""<div class="onboarding-box"><h4>👋 Olá, Arthur!</h4><p>Sou seu <b>Mentor Jurídico</b>. Dúvidas, peças ou jurisprudência? Estou conectado às IAs mais robustas do mercado e indexado ao Google para fatos recentes.</p></div>""", unsafe_allow_html=True)
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"], avatar="🧑‍⚖️" if msg["role"] == "user" else "🤖"): 
                st.markdown(msg["content"])
        if p := st.chat_input("Digite..."):
            st.session_state.chat_history.append({"role": "user", "content": p})
            with st.chat_message("user", avatar="🧑‍⚖️"): 
                st.write(p)
            with st.chat_message("assistant", avatar="🤖"):
                with st.spinner("Analisando com IA de última geração..."):
                    history = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.chat_history[-6:]])
                    res = call_gemini("Você é um Advogado Sênior e Assessor Jurídico de alto nível. Forneça respostas técnicas embasadas e precisas.", history, use_search=True)
                    st.write(res)
                    st.session_state.chat_history.append({"role": "assistant", "content": res})
                    add_xp(5)

elif menu == "📝 Gere seu Contrato":
    if verificar_acesso_premium():
        st.title("📝 Gere seu Contrato")
        step = st.session_state.contract_step
        
        if step == 1:
            st.markdown("""<div class="onboarding-box"><b>Crie minutas perfeitas.</b><br>Escolha o tipo, informe as partes e a IA redige.</div>""", unsafe_allow_html=True)
        c1, c2, c3 = st.columns([1,1,1])
        c1.markdown(f"**1. Dados** {'✅' if step > 1 else '🟦'}")
        c2.markdown(f"**2. Minuta** {'✅' if step > 2 else ('🟦' if step==2 else '⬜')}")
        c3.markdown(f"**3. Baixar** {'✅' if step > 3 else ('🟦' if step==3 else '⬜')}")
        st.progress(int(step/3 * 100))

        if step == 1:
            with st.container(border=True):
                tipo = st.selectbox("Modelo:", ["Prestação de Serviços", "Locação de Imóvel", "Compra e Venda Imóvel", "Compra e Venda Veículo", "Outro"])
                partes = st.text_area("Partes")
                objeto = st.text_area("Objeto")
                if st.button("Gerar Minuta ➔", type="primary", use_container_width=True):
                    if partes and objeto:
                        with st.spinner("Redigindo com IA estrutural..."):
                            lei = "Lei do Inquilinato" if "Locação" in tipo else "Código Civil"
                            prompt = f"Crie contrato de {tipo}. Base: {lei}. Partes: {partes}. Objeto: {objeto}. JSON: {{'clauses': [{{'titulo': '...', 'conteudo': '...'}}]}}"
                            res = call_gemini("JSON only.", prompt, json_mode=True)
                            data = extract_json_surgical(res)
                            if data and 'clauses' in data:
                                st.session_state.contract_meta = {"tipo": tipo, "partes": partes, "objeto": objeto}
                                st.session_state.contract_clauses = data['clauses']
                                st.session_state.contract_step = 2
                                add_xp(25)
                                st.rerun()
                            else: 
                                st.error("Erro ao gerar.")
        elif step == 2:
            st.header("📑 Revisão")
            if st.button("➕ Cláusula"): 
                st.session_state.contract_clauses.append({"titulo":"Nova","conteudo":"..."})
                st.rerun()
            to_remove = []
            for i, c in enumerate(st.session_state.contract_clauses):
                with st.expander(f"{i+1}. {c.get('titulo')}", expanded=False):
                    nt = st.text_input("T",c['titulo'],key=f"t{i}")
                    nc = st.text_area("C",c['conteudo'],key=f"c{i}")
                    st.session_state.contract_clauses[i] = {"titulo":nt,"conteudo":nc}
                    if st.button("🗑️",key=f"d{i}"): 
                        to_remove.append(i)
            if to_remove:
                for i in sorted(to_remove, reverse=True): 
                    del st.session_state.contract_clauses[i]
                st.rerun()
            c1,c2=st.columns([1,2])
            if c1.button("⬅️"): 
                st.session_state.contract_step=1
                st.rerun()
            if c2.button("Finalizar ➔",type="primary",use_container_width=True): 
                st.session_state.contract_step=3
                st.rerun()
        elif step == 3:
            st.header("✅ Pronto")
            docx_bytes = create_contract_docx(st.session_state.contract_clauses, st.session_state.contract_meta)
            if docx_bytes: 
                st.download_button("💾 Baixar DOCX", docx_bytes, "Contrato.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", use_container_width=True)
            if st.button("✏️ Editar"): 
                st.session_state.contract_step=2
                st.rerun()

elif menu == "🎯 Mestre dos Editais":
    if verificar_acesso_premium():
        st.title("🎯 Mestre dos Editais")
        
        if not st.session_state.edital_text:
            st.markdown("""
            <div class="onboarding-box">
                <h4>🚀 Professor de Edital</h4>
                <p>Transforme PDF em simulador de prova usando inteligência contextual profunda.</p>
                <ul><li>Carregue o Edital > Gere questões técnicas > Aprenda.</li></ul>
            </div>
            """, unsafe_allow_html=True)

        def gerar_turbo(dificuldade, foco):
            st.session_state.quiz_data = None
            st.session_state.quiz_show_answer = False
            with st.spinner(f"⚡ Criando questão ({dificuldade})..."):
                tema = f"FOCO: {foco}." if foco else "Tema aleatório."
                txt = st.session_state.edital_text
                
                prompt = f"""
                Role: Banca Examinadora de Concursos Públicos.
                TASK: Criar questão técnica inédita de múltipla escolha.
                CRITICAL: IGNORE TOTALMENTE datas, inscrições, taxas, isenções, locais de prova e vagas.
                SOURCE: Busque APENAS nos ANEXOS de 'CONTEÚDO PROGRAMÁTICO' ou 'CONHECIMENTOS ESPECÍFICOS'.
                {tema} Nível: {dificuldade}.
                JSON Output: {{'materia':'...','enunciado':'...','alternativas':{{'A':'...','B':'...','C':'...','D':'...'}},'correta':'A','explicacao':'...'}}
                """
                
                res = call_gemini("JSON Only.", f"{prompt}\nEDITAL:\n{txt}", json_mode=True)
                if "Limite de velocidade" in res:
                    st.error(res)
                else:
                    data = extract_json_surgical(res)
                    if data: 
                        st.session_state.quiz_data = data
                    else: 
                        st.error("Erro rápido.")

        if not st.session_state.edital_text:
            f = st.file_uploader("Upload PDF", type=["pdf"])
            if f and f.name != st.session_state.edital_filename:
                with st.spinner("Lendo (pode demorar um pouco se for grande)..."):
                    txt = read_pdf_safe(f)
                    if txt: 
                        st.session_state.edital_text=txt
                        st.session_state.edital_filename=f.name
                        st.rerun()
                    else: 
                        st.error("PDF sem texto.")
        else:
            c1, c2 = st.columns([3, 1])
            c1.success(f"📂 **{st.session_state.edital_filename}**")
            if c2.button("🗑️ Trocar", use_container_width=True): 
                st.session_state.edital_text=""
                st.rerun()
            st.markdown("---")
            cc, ca = st.columns([2, 1])
            with cc:
                diff = st.select_slider("Nível:", ["Fácil", "Médio", "Difícil", "Pesadelo"], value="Difícil")
                foco = st.text_input("Foco:", placeholder="Ex: Penal")
            with ca:
                st.write(""); st.write("")
                if st.button("🔥 GERAR", type="primary", use_container_width=True): 
                    gerar_turbo(diff, foco)
                    st.rerun()

            if st.session_state.quiz_data:
                q = st.session_state.quiz_data
                st.markdown(f"### 📚 {q.get('materia','Geral')}")
                st.info(q['enunciado'])
                opts = q['alternativas']
                if not st.session_state.quiz_show_answer:
                    c1,c2 = st.columns(2)
                    if c1.button(f"A) {opts['A']}", use_container_width=True): 
                        st.session_state.user_choice="A"
                        st.session_state.quiz_show_answer=True
                        st.rerun()
                    if c2.button(f"B) {opts['B']}", use_container_width=True): 
                        st.session_state.user_choice="B"
                        st.session_state.quiz_show_answer=True
                        st.rerun()
                    if c1.button(f"C) {opts['C']}", use_container_width=True): 
                        st.session_state.user_choice="C"
                        st.session_state.quiz_show_answer=True
                        st.rerun()
                    if c2.button(f"D) {opts['D']}", use_container_width=True): 
                        st.session_state.user_choice="D"
                        st.session_state.quiz_show_answer=True
                        st.rerun()
                else:
                    u, c = st.session_state.user_choice, q['correta']
                    for l,t in opts.items():
                        icon = "✅" if l==c else ("❌" if l==u else "⬜")
                        st.write(f"{icon} **{l})** {t}")
                    if u==c: 
                        st.success("Acertou!")
                        add_xp(50)
                    else: 
                        st.error(f"Errou. Correta: {c}")
                    st.write(f"**Explicação:** {q['explicacao']}")
                    
                    q_text = f"MATÉRIA: {q['materia']}\n\nQUESTÃO:\n{q['enunciado']}\n\nA) {opts['A']}\nB) {opts['B']}\nC) {opts['C']}\nD) {opts['D']}\n\nRESPOSTA: {q['correta']}\n\nCOMENTÁRIO:\n{q['explicacao']}"
                    docx_q = create_generic_docx(q_text, "Questão de Concurso")
                    if docx_q:
                        st.download_button("💾 Baixar Questão (Word)", docx_q, "Questao_Carmelio.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    if st.button("➡️ Próxima", type="primary"): 
                        gerar_turbo(diff, foco)
                        st.rerun()

elif menu == "🏢 Cartório OCR":
    if verificar_acesso_premium():
        st.title("🏢 Cartório OCR (Digitalizador)")
        st.markdown("""
        <div class="onboarding-box">
            <h4>📸 Do Papel para o Digital</h4>
            <p>Digitalize folhas e transcrições com visão computacional de alta fidelidade.</p>
            <ul><li><b>Envie:</b> Foto da página do livro.</li><li><b>Receba:</b> Texto transcrito para Certidão de Inteiro Teor.</li></ul>
        </div>
        """, unsafe_allow_html=True)
        
        img_file = st.file_uploader("Foto do Livro/Documento", type=["png", "jpg", "jpeg"])
        if img_file:
            image = Image.open(img_file)
            st.image(image, caption="Imagem Carregada", use_container_width=True)
            if st.button("🔍 Extrair Texto", type="primary"):
                with st.spinner("Processando OCR com IA Multimodal..."):
                    res = call_gemini("Especialista em OCR cartorial. Transcreva TODO o texto com precisão total, mantendo nomes e datas.", "Transcreva.", image=image)
                    if "Limite de velocidade" in res:
                        st.error(res)
                    else:
                        st.session_state.ocr_text = res
                        add_xp(30)
        
        if st.session_state.ocr_text:
            st.text_area("Texto Extraído:", st.session_state.ocr_text, height=400)
            docx_ocr = create_generic_docx(st.session_state.ocr_text, "Transcrição de Livro")
            if docx_ocr:
                st.download_button("💾 Baixar Texto em Word", docx_ocr, "Certidao_Inteiro_Teor.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")

elif menu == "🎙️ Transcrição":
    if verificar_acesso_premium():
        st.title("🎙️ Transcrição de Áudio Real")
        st.markdown("""<div class="onboarding-box"><h4>🗣️ Voz para Texto Inteligente</h4><p>Envie o arquivo e a inteligência artificial fará a transcrição completa e a estruturação lógica do conteúdo.</p></div>""", unsafe_allow_html=True)
        
        audio_file = st.file_uploader("Arquivo de Áudio", type=["mp3", "wav", "m4a", "ogg"])
        if audio_file:
            st.audio(audio_file)
            
            if st.button("📝 Transcrever Áudio Real", type="primary"):
                with st.spinner("Enviando arquivo para processamento neural..."):
                    mime = "audio/mp3" if audio_file.name.endswith("mp3") else "audio/wav"
                    if audio_file.name.endswith("m4a"): mime = "audio/m4a"
                    if audio_file.name.endswith("ogg"): mime = "audio/ogg"
                    
                    raw_data = audio_file.getvalue()
                    
                    system_instruction = "Você é um estenógrafo técnico e assessor jurídico especialista. Transcreva fielmente o áudio fornecido, organizando o texto em parágrafos coerentes. Se houver termos jurídicos, garanta a grafia correta."
                    
                    res = call_gemini(system_instruction, "Transcreva este áudio na íntegra.", audio_bytes=raw_data, audio_mime=mime)
                    
                    if "Limite de velocidade" in res:
                        st.error(res)
                    else:
                        st.session_state.audio_text = res
                        add_xp(40)
        
        if st.session_state.audio_text:
            st.text_area("Resultado da Transcrição Real:", st.session_state.audio_text, height=300)
            docx_audio = create_generic_docx(st.session_state.audio_text, "Transcrição Judicial Inteligente")
            if docx_audio:
               st.download_button("💾 Baixar Transcrição (Word)", docx_audio, "Transcricao_Real.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")
